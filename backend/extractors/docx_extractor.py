"""Word belgelerinden metin çıkarır (paragraflar + tablolar)."""

from docx import Document

from .base import ExtractionError, ExtractionResult


def extract(path: str) -> ExtractionResult:
    try:
        doc = Document(path)
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        raw_text = "\n".join(parts)
        return ExtractionResult(
            raw_text=raw_text,
            metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
        )
    except Exception as e:
        raise ExtractionError(f"DOCX işlenemedi: {e}") from e
